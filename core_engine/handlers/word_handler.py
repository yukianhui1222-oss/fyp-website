"""
word_handler.py
───────────────
Word (.docx) 文档处理器，基于 python-docx。

处理策略：
  - 遍历所有段落和表格，提取纯文本。
  - 收集 Document 中的内嵌图片（通过 Document.part.rels），
    送入 OCR 处理器进行识别。
"""

import io
import logging
import os
import tempfile
from typing import Any, Callable, Dict, List, Optional

from PIL import Image

logger = logging.getLogger(__name__)


class WordHandler:
    """
    Word 文档处理器。

    Parameters
    ----------
    ocr_processor : OCRProcessor
        已初始化的 OCR 处理器实例（由 DocumentParser 传入）
    """

    def __init__(self, ocr_processor: Any) -> None:
        self.ocr = ocr_processor

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """Use Word OLE automation or LibreOffice to convert .doc to .docx"""
        # Try win32com first on Windows
        if os.name == 'nt':
            try:
                import win32com.client
                logger.info("Using Word OLE automation for conversion: %s", doc_path)
                abs_doc_path = os.path.abspath(doc_path).replace('/', '\\')
                temp_dir = tempfile.gettempdir()
                base_name = os.path.basename(doc_path)
                tmp_docx_path = os.path.join(temp_dir, f"conv_{base_name}x")
                abs_docx_path = os.path.abspath(tmp_docx_path).replace('/', '\\')

                # Ensure CoInitialize
                import pythoncom
                pythoncom.CoInitialize()

                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                # Open: FileName, ReadOnly
                document = word.Documents.Open(abs_doc_path, ReadOnly=True)
                document.SaveAs2(abs_docx_path, 16) # 16 = wdFormatDocumentDefault (.docx)
                document.Close()
                word.Quit()
                return abs_docx_path
            except Exception as e:
                logger.warning("Word OLE automation failed or pywin32 missing: %s. Trying LibreOffice...", e)
        
        # Try LibreOffice conversion
        try:
            import subprocess
            import shutil
            soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice_path and os.name == 'nt':
                standard_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                for p in standard_paths:
                    if os.path.exists(p):
                        soffice_path = p
                        break
            
            if not soffice_path:
                raise RuntimeError("LibreOffice/soffice executable not found in system PATH.")
                
            temp_dir = tempfile.gettempdir()
            cmd = [soffice_path, "--headless", "--convert-to", "docx", doc_path, "--outdir", temp_dir]
            logger.info("Running LibreOffice command: %s", " ".join(cmd))
            
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice failed with code {result.returncode}: {result.stderr}")
                
            base_name = os.path.basename(doc_path)
            root_name, _ = os.path.splitext(base_name)
            output_filename = f"{root_name}.docx"
            output_path = os.path.join(temp_dir, output_filename)
            
            if not os.path.exists(output_path):
                raise RuntimeError("Converted file not found at expected path.")
                
            return output_path
        except Exception as e:
            logger.error("All conversion methods failed: %s", e)
            raise RuntimeError(
                "System missing conversion libraries. Please ensure LibreOffice is installed (listed in packages.txt on Streamlit Cloud) or upload .docx files directly. Error details: " + str(e)
            )

    # ------------------------------------------------------------------
    def process(self, file_path: str, progress_callback: Optional[Callable] = None) -> List[Dict[str, Any]]:
        """
        Parses Word files and returns a structured list of results ordered by page (Word is actually streaming, simulated here as a single page or blocks).

        Parameters
        ----------
        file_path : str
            Absolute path to the Word file (.docx)
        progress_callback : Optional[callable]
            Progress callback

        Returns
        -------
        list[dict]
            Single-element list (page_num=1) containing all text blocks and OCR blocks.
        """
        is_legacy = file_path.lower().endswith(".doc")
        temp_docx = None

        if is_legacy:
            try:
                temp_docx = self._convert_doc_to_docx(file_path)
                file_path = temp_docx
            except Exception as e:
                return [
                    {
                        "page_num": 1,
                        "page_type": "error",
                        "text_blocks": [
                            {
                                "text": f"⚠️ Conversion Error: Unable to convert .doc to .docx. {e}",
                                "bbox": None,
                                "source": "error",
                                "confidence": 0.0,
                            }
                        ],
                        "table_html": None,
                    }
                ]

        import docx  # python-docx

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.error("Unable to open Word file: %s", e)
            return [
                {
                    "page_num": 1,
                    "page_type": "error",
                    "text_blocks": [
                        {
                            "text": f"⚠️ Unable to parse this Word file. The file may be corrupted or in an unsupported format. Error details: {e}",
                            "bbox": None,
                            "source": "error",
                            "confidence": 0.0,
                        }
                    ],
                    "table_html": None,
                }
            ]

        logger.info("Starting Word document analysis: %s", file_path)

        native_blocks: List[Dict[str, Any]] = []
        ocr_images: List[Image.Image] = []

        # ── 提取段落文本 ──────────────────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                native_blocks.append(
                    {
                        "text": text,
                        "bbox": None,       # Word 段落无绝对坐标
                        "source": "native",
                        "confidence": 1.0,
                    }
                )

        # ── 提取表格文本 ──────────────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        native_blocks.append(
                            {
                                "text": cell_text,
                                "bbox": None,
                                "source": "native",
                                "confidence": 1.0,
                            }
                        )

        # ── 提取内嵌图片 ──────────────────────────────────────────────────
        for rel in doc.part.rels.values():
            # 通过 MIME 类型过滤图片关系
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    ocr_images.append(img)
                except Exception as exc:  # noqa: BLE001
                    logger.warning("Failed to extract Word image: %s", exc)

        # ── 并发 OCR ──────────────────────────────────────────────────────
        ocr_results = self.ocr.process_images(ocr_images, progress_callback=progress_callback) if ocr_images else []
        ocr_blocks: List[Dict[str, Any]] = []
        for res in ocr_results:
            if res.get("error"):
                logger.warning("Word image OCR failed (index=%d): %s", res["index"], res["error"])
            ocr_blocks.extend(res["blocks"])

        # ── 组装结构 ──────────────────────────────────────────────────────
        all_blocks = native_blocks + ocr_blocks
        logger.info(
            "Word analysis complete | Native text blocks=%d | OCR images=%d",
            len(native_blocks), len(ocr_images)
        )

        # ── 清理临时文件 ──────────────────────────────────────────────
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
                logger.info("Cleaned up temporary conversion file: %s", temp_docx)
            except Exception as e:
                logger.warning("Failed to clean up temporary file: %s", e)

        return [
            {
                "page_num": 1,
                "page_type": "native",
                "text_blocks": all_blocks,
                "table_html": None,
            }
        ]
